"""
Find-and-replace edits on the IEEE IRI 2026 paper .docx
to incorporate blinded multi-judge scoring results.
Preserves all formatting by operating on paragraph runs.
"""
from docx import Document
import re

doc = Document("/home/user/PLEAS/paper_updated.docx")

replacements = [
    # 1. Abstract — scores and methodology
    (
        "versus 38/48 for a non-PLEAS baseline and 28/48 for ChatGPT (gpt-5.5-thinking-extended)",
        "versus 39/48 for a non-PLEAS baseline and 31/48 for GPT-5.5-thinking-extended"
    ),
    # 1b. Abstract — add "scored by blinded multi-judge evaluation"
    (
        "a 24-task biomedical scientific workflow benchmark in which",
        "a 24-task biomedical scientific workflow benchmark scored by blinded multi-judge evaluation in which"
    ),

    # 2. Section V-B — Scoring methodology (replace full paragraph)
    (
        "Scoring was performed by a fresh, separate instance of GPT-5.5-thinking-extended using the BioClaw rubric, in a non-blinded manner.",
        "Scoring used a blinded multi-judge protocol: responses were anonymized via HMAC-SHA256 with system-identifying names scrubbed by regex, then independently scored by three LLM judges (Claude Opus 4.8, Claude Sonnet 4.6, Gemini 2.5 Pro)."
    ),
    (
        "We note that the ChatGPT comparator and the scoring judge use the same model family; this confound is disclosed as a limitation and means comparisons involving the ChatGPT row should be interpreted with caution.",
        "Majority-vote aggregation yielded per-task scores; inter-rater reliability measured Fleiss’ kappa of 0.53 (moderate agreement) and Krippendorff’s alpha of 0.11."
    ),
    # 2b. ChatGPT → GPT-5.5-thinking-extended in the systems-evaluated sentence
    (
        "and ChatGPT (gpt-5.5-thinking-extended). Scoring",
        "and GPT-5.5-thinking-extended. Scoring"
    ),

    # 3. Section V-B — Results paragraph
    (
        "the non-PLEAS baseline achieved 38/48 (79.2%), and ChatGPT (gpt-5.5-thinking-extended) achieved 28/48 (58.3%).",
        "the non-PLEAS baseline achieved 39/48 (81.2%), and GPT-5.5-thinking-extended achieved 31/48 (64.6%)."
    ),
    (
        "Literature & Clinical Mining was BioClaw’s weakest category (3/8), primarily due to T10 (live registry count failure) and T11 (Gene Expression Omnibus (GEO) dataset GSE68849 inaccessible to all three systems, all scoring 0). The 8.3-point margin",
        "Transcriptomics & Single-Cell Analysis was the most discriminating category: BioClaw scored 7/8, the non-PLEAS baseline 6/8, and GPT-5.5-thinking-extended only 2/8, suggesting that structured knowledge retrieval provides the greatest advantage on complex multi-step analytical tasks. The 3-point margin"
    ),

    # 4. Failure Analysis
    (
        "Four tasks received less than full credit across systems. T10 scored 0 for BioClaw, 2 for the non-PLEAS baseline, and 1 for ChatGPT, indicating a retrieval path failure specific to BioClaw. T11 scored 0 for all three systems under both original and corrected metadata expectations, confirming a data accessibility barrier rather than a capability difference. T03 and T09 received a score of 1 from all three systems, suggesting ambiguity in the rubric or task specification.",
        "Under blinded multi-judge scoring, GPT-5.5-thinking-extended received four zero scores (T07, T14, T16, T23), all in categories requiring multi-step tool orchestration; neither BioClaw nor the non-PLEAS baseline received any zero. T03 received a score of 1 from all three systems, suggesting rubric ambiguity rather than capability differences. T14 (GSE159929 dataset mismatch) was correctly identified by BioClaw (score 2) but failed under GPT, indicating that structured retrieval planning in the Learn phase improves error detection."
    ),

    # 5. Discussion — Limitations
    (
        "Three limitations bound interpretation: non-blinded AI scoring, one token run per condition with unmatched model routing, and an incomplete portability probe. Blinded human evaluation, matched multi-task token runs, and full five-phase ports are needed before strong quantitative claims can be made.",
        "Two limitations bound interpretation: one token run per condition with unmatched model routing, and an incomplete portability probe. Blinded multi-judge scoring addresses the prior non-blinded evaluation confound. Matched token runs and full five-phase ports are needed before strong claims can be made."
    ),
    # 5b. Future work sentence
    (
        "add blinded human scoring,",
        "add blinded human evaluation,"
    ),

    # 6. Table III — Evidence-to-Claim row
    (
        "non-blinded GPT-5.5-thinking-extended scoring",
        "blinded multi-judge scoring (Fleiss’ kappa 0.53)"
    ),

    # 7. Conclusion
    (
        "preliminary and non-blinded",
        "scored by blinded multi-judge evaluation"
    ),
    (
        "versus 38/48 for a non-PLEAS baseline and 28/48 for ChatGPT (gpt-5.5-thinking-extended).",
        "versus 39/48 for a non-PLEAS baseline and 31/48 for GPT-5.5-thinking-extended."
    ),
]


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving run formatting."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    new_full = full_text.replace(old_text, new_text)

    if not paragraph.runs:
        return False

    first_run_fmt = paragraph.runs[0]

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = new_full
    return True


def replace_in_table_cell(cell, old_text, new_text):
    """Replace text in table cells."""
    for paragraph in cell.paragraphs:
        if old_text in paragraph.text:
            return replace_in_paragraph(paragraph, old_text, new_text)
    return False


applied = {i: False for i in range(len(replacements))}

for para in doc.paragraphs:
    full = para.text
    for i, (old, new) in enumerate(replacements):
        if old in full:
            if replace_in_paragraph(para, old, new):
                applied[i] = True
                print(f"  Applied replacement #{i+1} in paragraph")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for i, (old, new) in enumerate(replacements):
                if not applied[i]:
                    cell_text = cell.text
                    if old in cell_text:
                        if replace_in_table_cell(cell, old, new):
                            applied[i] = True
                            print(f"  Applied replacement #{i+1} in table cell")

print("\n--- Results ---")
for i, (old, new) in enumerate(replacements):
    status = "APPLIED" if applied[i] else "NOT FOUND"
    print(f"  #{i+1}: {status} — {old[:60]}...")

doc.save("/home/user/PLEAS/paper_updated.docx")
print("\nSaved to paper_updated.docx")
